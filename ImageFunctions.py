import numpy as np
import numba as nb
from PIL import Image, ImageSequence
from pathlib import Path
from tqdm import tqdm
from docx import Document
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import ImageFormatter
import cv2
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# 92 it/s, #parallel 110 it/s
@nb.njit(parallel=True)
def remove_background_from_numpy_image(image_array, gray_array, min_gray, max_gray):
    x_range, y_range = gray_array.shape
    for x in nb.prange(x_range):
        for y in nb.prange(y_range):
            if min_gray <= gray_array[x, y] <= max_gray:
                image_array[x, y, 3] = 0
    return image_array


# 45 it/s
# def remove_background_from_numpy_image(image_array, gray_array, min_gray, max_gray):
#     image_array[(gray_array >= min_gray) & (gray_array <= max_gray)] = 0
#     return image_array


def remove_background_from_image(image: Image, min_gray, max_gray):
    rgba_image = image.convert("RGBA")
    img_array = np.array(rgba_image)
    gray_array = np.array(rgba_image.convert("L"))
    img_array = remove_background_from_numpy_image(img_array, gray_array, min_gray, max_gray)
    return Image.fromarray(img_array, 'RGBA')


def remove_background_from_gif(gif_path, output_path, lower_bg_color, upper_bg_color, duration):
    gif = Image.open(gif_path)
    frames = []
    for frame in tqdm(ImageSequence.Iterator(gif)):
        frames.append(remove_background_from_image(frame, lower_bg_color, upper_bg_color))
    frames[0].save(output_path, save_all=True, append_images=frames[1:],
                   duration=duration, loop=0, disposal=2)


def extract_background_and_save_frames(gif_path, output_dir, lower_bg_color, upper_bg_color):
    gif = Image.open(gif_path)
    for index, frame in tqdm(enumerate(ImageSequence.Iterator(gif))):
        frame_transparent = remove_background_from_image(frame, lower_bg_color, upper_bg_color)
        frame_transparent.save(output_dir / f'frame_{str(index).zfill(5)}.png', quality=95)


def create_gif(images, target, name, fps, add_reverse=False, loop=1):
    fourcc = cv2.VideoWriter_fourcc(*'H264')
    all_images = list((images / name).glob('*.jpg'))
    height, width, _ = cv2.imread(str(all_images[0])).shape
    video = cv2.VideoWriter(str(target / name) + '.gif', fourcc, fps, (width, height))
    for i in range(loop):
        for index, img in tqdm(enumerate(all_images)):
            if i > 0 and index == 0:
                continue
            video.write(cv2.imread(str(img)))
        if add_reverse:
            for img in tqdm((all_images[-2::-1])):
                video.write(cv2.imread(str(img)))
    cv2.destroyAllWindows()
    video.release()


def add_white_background(image: Image, color='white') -> Image:
    size = max(image.size)
    x = (size - image.width) // 2
    y = (size - image.height) // 2
    new_image = Image.new('RGB', (size, size), color=color)
    new_image.paste(image, (x, y))
    return new_image


def crop_to_ratio_4_5(image: Image, color='white') -> Image:
    output_size = max(image.size)
    if image.width != image.height:
        image = add_white_background(image, color=color)
    new_width = int(image.width * 0.8)
    left_margin = (output_size - new_width) // 2
    right_margin = output_size - left_margin
    return image.crop((left_margin, 0, right_margin, output_size))


def reshape_all_images_in_directory(directory: Path, target: Path,
                                    image_format='.jpg',
                                    color='white', square=False):
    for path in directory.glob(f'*{image_format}'):
        print(path)
        image = Image.open(path)
        if image.width > image.height or square:
            new_image = add_white_background(image, color=color)
        else:
            new_image = crop_to_ratio_4_5(image, color=color)
        new_image.save(target / path.name, quality=95)


def read_images_and_save_as_docx(images_path: Path):
    doc_eng = Document()
    doc_pol = Document()
    for image_path in images_path.glob('*.jpg'):
        print(image_path.name)
        image = Image.open(image_path)
        for doc, lang in zip([doc_eng, doc_pol], ['eng', 'pol']):
            text = pytesseract.image_to_string(image, lang=lang)
            doc.add_paragraph(text)
            doc.add_page_break()
    doc_eng.save('./ImagesToTesseract/eng.docx')
    doc_pol.save('./ImagesToTesseract/pol.docx')


def convert_code_to_formatted_font(code, save_path, transparent_bg: tuple = None, **format_params):
    """transparent_bg is a tuple of (min_gray, max_gray)"""
    formatter = ImageFormatter(**format_params)
    with open(save_path, 'wb') as file:
        file.write(highlight(code.encode(), PythonLexer(), formatter))
    if transparent_bg is not None:
        image = Image.open(save_path)
        image = remove_background_from_image(image, *transparent_bg)
        image.save(save_path, quality=95)


if __name__ == '__main__':
    gif_path = Path(
        r'C:\Users\thejg\Desktop\Studia\BigDataAnalytics2024\Seminars\BoidsFlockingPresentation\graphics\slide3gif.gif')
    output_path = Path(
        r'C:\Users\thejg\Desktop\Programming\MediaEditor\outputs\GIFS\RemovingBackground\test.gif')
    gray_scale = (0, 50)
    # img = Image.open(gif_path)
    remove_background_from_gif(gif_path, output_path, *gray_scale, 1000 / 30)
    # img = remove_background_from_image(img, *gray_scale)
    # img.save(output_path, quality=95)
